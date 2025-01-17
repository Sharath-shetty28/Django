from django.shortcuts import render
from django.http import HttpResponse
from PyPDF2 import PdfMerger


def index(request):
    return render(request, 'merge/index.html')

def merge_pdfs(request):
    if request.method == "POST":
        files = request.FILES.getlist('pdfs')
        if not files:
            return HttpResponse("No files uploaded.", status=400)

        merger = PdfMerger()
        first_file_name = files[0].name.split('.')[0]  # Extract the name of the first file (without extension)

        for file in files:
            merger.append(file)

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{first_file_name}_merged.pdf"'
        merger.write(response)
        merger.close()

        return response

    return render(request, 'merge/merge_pdfs.html')

from django.http import JsonResponse, FileResponse
from django.shortcuts import render
from docx import Document
from io import BytesIO


def merge_words(request):
    if request.method == 'POST':
        files = request.FILES.getlist('words')

        # Ensure at least two files are uploaded
        if len(files) < 2:
            return JsonResponse({'error': 'Please upload at least two Word files.'}, status=400)

        # Create a new Word document to merge into
        merged_document = Document()

        for file in files:
            if file.name.endswith('.docx'):
                doc = Document(file)
                for paragraph in doc.paragraphs:
                    merged_document.add_paragraph(paragraph.text)
            else:
                return JsonResponse({'error': f'{file.name} is not a valid Word document.'}, status=400)

        # Save merged document to a BytesIO object
        merged_file = BytesIO()
        merged_document.save(merged_file)
        merged_file.seek(0)

        # Generate merged file name
        first_file_name = files[0].name.split('.')[0]
        merged_file_name = f"{first_file_name}_merged.docx"

        # Return the merged file as a response
        response = FileResponse(merged_file, as_attachment=True, filename=merged_file_name)
        return response

    return render(request, 'merge/merge_words.html')
    
from django.shortcuts import render
from django.http import HttpResponse, JsonResponse
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO

def lock_pdf(request):
    if request.method == 'POST':
        # Get the uploaded file and password
        pdf_file = request.FILES.get('pdf_file')
        password = request.POST.get('password')

        # Validate inputs
        if not pdf_file or not password:
            return JsonResponse({'error': 'PDF file and password are required.'}, status=400)

        try:
            # Read the PDF file
            pdf_reader = PdfReader(pdf_file)

            # Create a new PDF writer
            pdf_writer = PdfWriter()

            # Copy pages from reader to writer
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)

            # Encrypt the PDF with the given password
            pdf_writer.encrypt(password)

            # Save the locked PDF to a BytesIO object
            output_pdf = BytesIO()
            pdf_writer.write(output_pdf)
            output_pdf.seek(0)

            # Return the locked PDF as a downloadable file
            response = HttpResponse(output_pdf, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="locked.pdf"'
            return response

        except Exception as e:
            return JsonResponse({'error': f'Failed to lock PDF: {str(e)}'}, status=500)

    # Render the upload form if GET request
    return render(request, 'security/lock_pdf.html')
